from docx import Document
from docx.shared import RGBColor, Pt
from common import SYMBOL_BIBLE_VERSE_CONTAINER as SBVC, SYMBOL_BIBLE_VERSE_START as SBVS, titles
from bs4.element import Tag
from file import save

class Writer:
    def __init__(self, title: str) -> None:
        self.title = title
        self.doc = Document()
        self.doc.add_heading(title, 0)
    
    def add_intro(self, text: Tag):
        title = 'Introduction'
        text = text.replace(title, '')
        self.add_title(title, 18)
        self.doc.add_paragraph(text)
        self.doc.add_page_break()

    def add_title(self, title: str, size: int):
         paragraph = self.doc.add_paragraph()
         run = paragraph.add_run(title)
         run.font.size = Pt(size)

    def add_day(self, day: Tag):
        title = day.find_all(lambda tag: tag.name == 'strong' and tag.get_text() in titles)[0].text
        self.add_title(title, 18)
        day = day.text\
            .replace(title, '')\
            .replace('\n\n\n\n\n\n\n\n\n\n\n\n', '')\
            .replace('\n\n\n\n\n\n', '')\
            .replace('\n\n\n', '')
        texts = day.split(SBVC)
        for text in texts:
            if text[:2] == SBVS:
                text = text.replace(SBVS, '')
                paragraph = self.doc.add_paragraph()
                run = paragraph.add_run(text)
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
            else:
                paragraph = self.doc.add_paragraph(text)
        self.doc.add_page_break()

    def add_week(self, week: list[Tag], title: str):
        self.add_title(title.upper(), 24)
        for day in week:
            self.add_day(day)

    def lesson(self, intro: str, weeks: list[Tag]):
        self.add_intro(intro)
        for week in weeks:
            self.add_week(week['week'], week['title'])
        self.doc.save(self.title + '.docx')

    