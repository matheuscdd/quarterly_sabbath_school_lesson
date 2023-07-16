from docx import Document
from docx.shared import RGBColor, Pt

doc = Document()

# Adiciona o título de acordo com seu nível
doc.add_heading('Title', 1)

# Adiciona paragráfo normal
p = doc.add_paragraph()


run = p.add_run("Este é um texto estilizado")
run.bold = True
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)

# quebra a página
doc.add_page_break()

# quando quebra a página precisa adicionar um outro parágrafo
p = doc.add_paragraph('Lorem lorem')

# p.add_run('This text is colorfull').color = 'red'

doc.save('text.docx')
