from docx import Document
from docx.shared import RGBColor, Pt
from common import (
    SYMBOL_BIBLE_VERSE_CONTAINER as SBVC,
    SYMBOL_BIBLE_VERSE_START as SBVS,
    TITLES,
)
from bs4.element import Tag
from io import BytesIO


class Writer:
    def __init__(self, title: str, intro: Tag, weeks: list[dict]) -> None:
        self.title = title
        self.doc = Document()
        self.doc.add_heading(title, 0)
        self.__intro(intro)
        [self.__week(week["week"], week["title"]) for week in weeks]

    def docx(self) -> BytesIO:
        buffer = BytesIO()
        self.doc.save(buffer)
        buffer.seek(0)
        return buffer

    def __title(self, title: str, size: int):
        paragraph = self.doc.add_paragraph()
        run = paragraph.add_run(title)
        run.font.size = Pt(size)

    def __intro(self, text: Tag):
        title = "Introduction"
        text = text.replace(title, "")
        self.__title(title, 18)
        self.doc.add_paragraph(text)
        self.doc.add_page_break()

    def __week(self, week: list[Tag], title: str):
        self.__title(title.upper(), 24)
        for day in week:
            self.__day(day)

    def __day(self, day: Tag):
        title = day.find_all(
            lambda tag: tag.name == "strong" and tag.get_text() in TITLES
        )[0].text
        self.__title(title, 18)
        day = (
            day.text.replace(title, "")
            .replace("\n\n\n\n\n\n\n\n\n\n\n\n", "")
            .replace("\n\n\n\n\n\n", "")
            .replace("\n\n\n", "")
            .replace(" n ", "")
        )
        texts = day.split(SBVC)
        for text in texts:
            if text[:2] == SBVS:
                text = text.replace(SBVS, "")
                paragraph = self.doc.add_paragraph()
                run = paragraph.add_run(text)
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
            else:
                paragraph = self.doc.add_paragraph(text)
        self.doc.add_page_break()
